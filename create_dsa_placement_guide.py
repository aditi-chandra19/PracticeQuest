from html import escape

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DOCX = "DSA_Placement_and_LeetCode_Guide.docx"
OUTPUT_HTML = "DSA_Placement_and_LeetCode_Guide.html"


CORE_SYLLABUS = [
    ("0. C++ and STL Basics", [
        "Input/output, functions, recursion basics",
        "Arrays, strings, vectors",
        "pair, stack, queue, deque, priority_queue",
        "set, unordered_set, map, unordered_map",
        "Sorting with STL and custom comparators",
    ]),
    ("1. Complexity and Problem-Solving Basics", [
        "Time complexity, space complexity",
        "Big-O, Omega, Theta",
        "Best, average, worst case",
        "Dry run, tracing, and constraints-based thinking",
    ]),
    ("2. Arrays", [
        "Traversal, insertion, deletion, update",
        "Prefix sum, suffix sum, difference array",
        "Kadane's algorithm, majority element, Dutch national flag",
        "Rotate array, merge intervals, missing and repeating numbers",
        "Subarray problems",
    ]),
    ("3. Strings", [
        "Basic operations, frequency counting",
        "Palindrome and anagram problems",
        "Pattern matching basics",
        "String hashing basics",
        "Subsequence vs substring",
        "Longest common prefix, compression, manipulation",
    ]),
    ("4. Searching", [
        "Linear search",
        "Binary search",
        "Binary search on answer",
        "Lower bound and upper bound",
        "Search in rotated sorted array, peak element, first and last occurrence",
    ]),
    ("5. Sorting", [
        "Bubble sort, selection sort, insertion sort",
        "Merge sort, quick sort, heap sort",
        "Counting sort, radix sort basics",
        "Stable vs unstable sorting",
    ]),
    ("6. Bit Manipulation", [
        "AND, OR, XOR, NOT",
        "Left and right shift",
        "Set, unset, toggle, and check bit",
        "Count set bits, power of 2, XOR tricks, bitmask basics",
    ]),
    ("7. Hashing", [
        "Hash map, hash set, frequency array",
        "Collision idea",
        "Applications in arrays and strings",
        "Two sum, longest consecutive sequence, subarray sum",
    ]),
    ("8. Two Pointers and Sliding Window", [
        "Opposite-direction pointers and same-direction pointers",
        "Fixed-size and variable-size window",
        "Longest substring and minimum window patterns",
        "Subarray condition problems",
    ]),
    ("9. Linked List", [
        "Singly, doubly, and circular linked list",
        "Reverse list, middle of list, detect cycle",
        "Remove nth node, merge two sorted lists, intersection point",
        "LRU concept basics",
    ]),
    ("10. Stack", [
        "Stack using array or linked list",
        "Infix, postfix, prefix basics",
        "Balanced parentheses, next greater element",
        "Monotonic stack, largest rectangle in histogram, stock span",
    ]),
    ("11. Queue", [
        "Queue implementation, circular queue, deque",
        "Monotonic queue basics",
        "BFS use of queue",
    ]),
    ("12. Recursion and Backtracking Basics", [
        "Recursive thinking, base case, recursion tree",
        "Subsequence generation, permutations, combinations, subsets",
        "N-Queens, Rat in a Maze, Sudoku",
        "Basic pruning",
    ]),
    ("13. Trees", [
        "Terminology and binary tree representation",
        "Inorder, preorder, postorder",
        "BFS and level order",
        "Height, size, diameter, balanced tree",
        "Left, right, top, and bottom view",
        "LCA in binary tree, path sum, serialize and deserialize basics",
    ]),
    ("14. Binary Search Tree", [
        "Insert, search, delete",
        "Floor, ceil",
        "Kth smallest and kth largest",
        "Validate BST, BST iterator, LCA in BST",
    ]),
    ("15. Heap and Priority Queue", [
        "Min heap, max heap, heapify",
        "Kth largest or smallest",
        "Top K elements",
        "Merge K sorted arrays or lists",
        "Median from data stream concept",
    ]),
    ("16. Greedy", [
        "Activity selection, fractional knapsack, job sequencing",
        "Huffman coding, interval scheduling, gas station type logic",
        "Greedy proof intuition",
    ]),
    ("17. Graphs", [
        "Graph representation",
        "BFS, DFS, connected components",
        "Cycle detection in undirected and directed graph",
        "Bipartite graph",
        "Topological sort and Kahn's algorithm",
        "Shortest path in unweighted graph, Dijkstra, Bellman-Ford, Floyd-Warshall basics",
        "Minimum spanning tree, Prim's, Kruskal's",
        "DSU in graph problems, SCC basics, grid BFS and DFS",
    ]),
    ("18. Dynamic Programming", [
        "Memoization, tabulation, state design",
        "1D DP, 2D DP, subsequence DP",
        "Knapsack pattern, LIS, LCS, edit distance, coin change",
        "Matrix path problems, partition DP basics, interval DP basics, stock DP basics",
    ]),
    ("19. Basic Mathematics for DSA", [
        "GCD and LCM",
        "Prime numbers and sieve of Eratosthenes",
        "Fast exponentiation",
        "Modular arithmetic",
        "Combinatorics basics",
        "Factors, divisors, overflow awareness",
    ]),
]


ADVANCED_SYLLABUS = [
    ("1. Advanced Trees", [
        "AVL tree, red-black tree",
        "B-tree and B+ tree concepts",
        "Trie",
        "Segment tree",
        "Lazy propagation",
        "Fenwick tree / BIT",
        "Order statistic tree concept",
    ]),
    ("2. Advanced Graphs", [
        "Disjoint Set Union in depth",
        "Strongly connected components",
        "Tarjan's algorithm, Kosaraju's algorithm",
        "Bridges and articulation points",
        "Lowest common ancestor in trees",
        "Binary lifting",
        "Shortest path in DAG",
        "0-1 BFS, multi-source BFS",
        "Dijkstra variants",
        "Network flow basics, max flow, min cut intro",
    ]),
    ("3. Advanced Dynamic Programming", [
        "DP on trees",
        "Digit DP",
        "Bitmask DP",
        "State compression DP",
        "Rerooting DP",
        "Divide and conquer DP optimization",
        "Knuth optimization concept",
        "Advanced stock DP and game DP",
    ]),
    ("4. Advanced String Algorithms", [
        "KMP, Rabin-Karp, Z algorithm",
        "Boyer-Moore basics",
        "Trie-based string problems",
        "Rolling hash",
        "Suffix array concept",
        "Suffix automaton concept",
        "Manacher's algorithm concept",
    ]),
    ("5. Advanced Range Query Structures", [
        "Sparse table",
        "Segment tree variants",
        "Fenwick tree variants",
        "Mo's algorithm concept",
    ]),
    ("6. Advanced Bit Manipulation", [
        "Bitmask enumeration",
        "Subset generation using bits",
        "XOR basis concept",
        "State compression",
    ]),
    ("7. Advanced Problem-Solving Paradigms", [
        "Meet in the middle",
        "Divide and conquer optimization",
        "Sweep line concept",
        "Coordinate compression",
        "Randomized algorithms basics",
    ]),
    ("8. Advanced Heaps and Priority Structures", [
        "Mergeable heaps concept",
        "Binomial heap concept",
        "Fibonacci heap concept",
        "Indexed priority queue concept",
    ]),
]


TOPIC_SHEET = [
    ("1. Arrays", {
        "Beginner": ["Two Sum", "Contains Duplicate", "Best Time to Buy and Sell Stock", "Maximum Subarray"],
        "Intermediate": ["Product of Array Except Self", "Merge Intervals", "3Sum", "Sort Colors"],
        "Advanced": ["First Missing Positive", "Trapping Rain Water"],
    }),
    ("2. Strings", {
        "Beginner": ["Valid Anagram", "Valid Palindrome", "Longest Common Prefix", "Implement strStr()"],
        "Intermediate": ["Group Anagrams", "Longest Substring Without Repeating Characters", "String Compression", "Palindrome Partitioning"],
        "Advanced": ["Minimum Window Substring", "Longest Palindromic Substring"],
    }),
    ("3. Hashing and Prefix Sum", {
        "Beginner": ["Ransom Note", "Isomorphic Strings", "Find the Difference", "Subarray Sum Equals K"],
        "Intermediate": ["Longest Consecutive Sequence", "Contiguous Array", "Top K Frequent Elements"],
        "Advanced": ["Prefix and Suffix Search", "Random Pick with Weight"],
    }),
    ("4. Binary Search", {
        "Beginner": ["Binary Search", "Search Insert Position", "First Bad Version", "Sqrt(x)"],
        "Intermediate": ["Search in Rotated Sorted Array", "Find Peak Element", "Koko Eating Bananas", "Capacity To Ship Packages Within D Days"],
        "Advanced": ["Median of Two Sorted Arrays", "Split Array Largest Sum"],
    }),
    ("5. Two Pointers and Sliding Window", {
        "Beginner": ["Valid Palindrome II", "Move Zeroes", "Squares of a Sorted Array"],
        "Intermediate": ["Container With Most Water", "3Sum", "Longest Repeating Character Replacement", "Permutation in String"],
        "Advanced": ["Minimum Window Substring", "Sliding Window Maximum"],
    }),
    ("6. Linked List", {
        "Beginner": ["Reverse Linked List", "Middle of the Linked List", "Merge Two Sorted Lists", "Linked List Cycle"],
        "Intermediate": ["Remove Nth Node From End of List", "Reorder List", "Intersection of Two Linked Lists", "Add Two Numbers"],
        "Advanced": ["Reverse Nodes in k-Group", "Copy List with Random Pointer"],
    }),
    ("7. Stack and Monotonic Stack", {
        "Beginner": ["Valid Parentheses", "Implement Stack using Queues", "Min Stack"],
        "Intermediate": ["Daily Temperatures", "Evaluate Reverse Polish Notation", "Asteroid Collision", "Decode String"],
        "Advanced": ["Largest Rectangle in Histogram", "Basic Calculator"],
    }),
    ("8. Queue and Heap", {
        "Beginner": ["Implement Queue using Stacks", "Last Stone Weight", "Kth Largest Element in a Stream"],
        "Intermediate": ["Top K Frequent Elements", "K Closest Points to Origin", "Find Median from Data Stream"],
        "Advanced": ["Merge k Sorted Lists", "Task Scheduler"],
    }),
    ("9. Recursion and Backtracking", {
        "Beginner": ["Subsets", "Permutations", "Generate Parentheses"],
        "Intermediate": ["Combination Sum", "Palindrome Partitioning", "Letter Combinations of a Phone Number", "Word Search"],
        "Advanced": ["N-Queens", "Sudoku Solver"],
    }),
    ("10. Trees", {
        "Beginner": ["Maximum Depth of Binary Tree", "Invert Binary Tree", "Same Tree", "Symmetric Tree"],
        "Intermediate": ["Binary Tree Level Order Traversal", "Diameter of Binary Tree", "Balanced Binary Tree", "Path Sum II"],
        "Advanced": ["Binary Tree Maximum Path Sum", "Serialize and Deserialize Binary Tree"],
    }),
    ("11. BST", {
        "Beginner": ["Search in a Binary Search Tree", "Minimum Absolute Difference in BST", "Range Sum of BST"],
        "Intermediate": ["Validate Binary Search Tree", "Kth Smallest Element in a BST", "Lowest Common Ancestor of a BST"],
        "Advanced": ["Delete Node in a BST", "Convert BST to Greater Tree"],
    }),
    ("12. Greedy", {
        "Beginner": ["Assign Cookies", "Can Place Flowers", "Lemonade Change"],
        "Intermediate": ["Jump Game", "Partition Labels", "Gas Station", "Non-overlapping Intervals"],
        "Advanced": ["Jump Game II", "Candy"],
    }),
    ("13. Graphs", {
        "Beginner": ["Find if Path Exists in Graph", "Flood Fill", "Number of Islands", "Max Area of Island"],
        "Intermediate": ["Clone Graph", "Course Schedule", "Rotting Oranges", "Pacific Atlantic Water Flow"],
        "Advanced": ["Word Ladder", "Alien Dictionary"],
    }),
    ("14. Advanced Graphs", {
        "Beginner": ["Network Delay Time"],
        "Intermediate": ["Cheapest Flights Within K Stops", "Min Cost to Connect All Points", "Redundant Connection"],
        "Advanced": ["Reconstruct Itinerary", "Critical Connections in a Network"],
    }),
    ("15. Dynamic Programming 1D", {
        "Beginner": ["Climbing Stairs", "Min Cost Climbing Stairs", "House Robber"],
        "Intermediate": ["House Robber II", "Coin Change", "Decode Ways", "Word Break"],
        "Advanced": ["Perfect Squares", "Integer Break"],
    }),
    ("16. Dynamic Programming 2D", {
        "Beginner": ["Unique Paths", "Minimum Path Sum"],
        "Intermediate": ["Longest Common Subsequence", "Edit Distance", "Partition Equal Subset Sum", "Target Sum"],
        "Advanced": ["Distinct Subsequences", "Burst Balloons"],
    }),
    ("17. Bit Manipulation", {
        "Beginner": ["Single Number", "Number of 1 Bits", "Counting Bits", "Power of Two"],
        "Intermediate": ["Reverse Bits", "Bitwise AND of Numbers Range", "Sum of Two Integers"],
        "Advanced": ["Maximum XOR of Two Numbers in an Array"],
    }),
    ("18. Trie and Advanced Structures", {
        "Beginner": ["Implement Trie (Prefix Tree)", "Design Add and Search Words Data Structure"],
        "Intermediate": ["Word Search II", "Replace Words"],
        "Advanced": ["Maximum XOR with Trie", "Range Sum Query - Mutable"],
    }),
]


BEST_ORDER = [
    "C++ STL and complexity",
    "Arrays, strings, hashing",
    "Binary search",
    "Sorting",
    "Two pointers and sliding window",
    "Linked list",
    "Stack and queue",
    "Recursion and backtracking",
    "Trees, BST, heap",
    "Graphs",
    "Greedy",
    "Dynamic programming",
    "Bit manipulation",
    "Math basics",
    "Advanced DSA topics",
]


KEY_FOCUS = [
    "Arrays",
    "Strings",
    "Hashing",
    "Binary search",
    "Two pointers",
    "Sliding window",
    "Stack",
    "Queue",
    "Linked list",
    "Trees",
    "BST",
    "Heap",
    "Graphs",
    "Greedy",
    "Dynamic programming",
    "Recursion and backtracking",
    "Bit manipulation",
]


PRACTICE_STAGES = [
    ("Phase 1: Foundation", "Arrays, strings, hashing, binary search"),
    ("Phase 2: Interview Core", "Two pointers, sliding window, linked list, stack, queue, heap"),
    ("Phase 3: Strong DSA", "Trees, BST, graphs, greedy, dynamic programming, math basics"),
    ("Phase 4: Advanced", "Trie, segment tree, DSU, Fenwick tree, SCC, bridges, articulation points, bitmask DP"),
]


TARGETS = [
    ("Placement-ready base", "120 to 150 solved problems with revision"),
    ("Strong company prep", "200 to 250 solved problems"),
    ("Very strong LeetCode base", "300 plus problems with revision and pattern notes"),
]


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margin_value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "DADCE0")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def ensure_style(document, name, style_type, base=None):
    styles = document.styles
    try:
        return styles[name]
    except KeyError:
        style = styles.add_style(name, style_type)
        if base is not None:
            style.base_style = styles[base]
        return style


def set_run_font(run, name="Arial"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def add_bullet(document, text, style_name="Body Bullet"):
    p = document.add_paragraph(style=style_name)
    p.add_run(text)


def add_numbered(document, text, style_name="Body Number"):
    p = document.add_paragraph(style=style_name)
    p.add_run(text)


def add_topic_level(document, label, items):
    p = document.add_paragraph(style="Level Label")
    run = p.add_run(f"{label}: ")
    run.bold = True
    set_run_font(run)
    p.add_run(", ".join(items))


def apply_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    title_style = document.styles["Title"]
    title_style.font.name = "Arial"
    title_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    title_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.space_after = Pt(10)

    subtitle_style = ensure_style(document, "Subtitle Custom", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    subtitle_style.font.name = "Arial"
    subtitle_style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    subtitle_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    subtitle_style.font.size = Pt(11)
    subtitle_style.font.color.rgb = RGBColor(85, 85, 85)
    subtitle_style.paragraph_format.space_after = Pt(8)

    for style_name, size, before, after in (
        ("Heading 1", 16, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    body_bullet = ensure_style(document, "Body Bullet", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    body_bullet.paragraph_format.left_indent = Inches(0.5)
    body_bullet.paragraph_format.first_line_indent = Inches(-0.25)
    body_bullet.paragraph_format.space_after = Pt(4)
    body_bullet.paragraph_format.line_spacing = 1.15

    body_number = ensure_style(document, "Body Number", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    body_number.paragraph_format.left_indent = Inches(0.5)
    body_number.paragraph_format.first_line_indent = Inches(-0.25)
    body_number.paragraph_format.space_after = Pt(4)
    body_number.paragraph_format.line_spacing = 1.15

    level_label = ensure_style(document, "Level Label", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    level_label.paragraph_format.space_after = Pt(4)
    level_label.paragraph_format.line_spacing = 1.15

    small_note = ensure_style(document, "Small Note", WD_STYLE_TYPE.PARAGRAPH, base="Normal")
    small_note.font.size = Pt(10)
    small_note.font.color.rgb = RGBColor(85, 85, 85)
    small_note.paragraph_format.space_after = Pt(6)


def add_footer(document):
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("DSA Placement and LeetCode Guide")
    set_run_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_intro(document):
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("DSA Placement and LeetCode Guide")

    subtitle = document.add_paragraph(style="Subtitle Custom")
    subtitle.add_run("A structured roadmap for placements, coding interviews, strong DSA fundamentals, and LeetCode practice.")

    note = document.add_paragraph(style="Small Note")
    note.add_run("This guide consolidates the full DSA syllabus, advanced DSA syllabus, best study order, and a topic-wise LeetCode progression from beginner to advanced.")


def add_syllabus_section(document, title_text, items):
    document.add_paragraph(title_text, style="Heading 1")
    for section_title, bullets in items:
        document.add_paragraph(section_title, style="Heading 2")
        for bullet in bullets:
            add_bullet(document, bullet)


def add_strategy_section(document):
    document.add_paragraph("Placement Strategy", style="Heading 1")

    document.add_paragraph("Best Order to Study", style="Heading 2")
    for item in BEST_ORDER:
        add_numbered(document, item)

    document.add_paragraph("Highest-Value Topics for Placements", style="Heading 2")
    for item in KEY_FOCUS:
        add_bullet(document, item)

    document.add_paragraph("Practical Phases", style="Heading 2")
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_borders(table)
    widths = [Inches(2.1), Inches(4.4)]
    hdr = table.rows[0].cells
    hdr[0].width = widths[0]
    hdr[1].width = widths[1]
    hdr[0].text = "Phase"
    hdr[1].text = "Focus"
    set_repeat_table_header(table.rows[0])

    for cell in hdr:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_run_font(run)
                run.bold = True

    for phase, focus in PRACTICE_STAGES:
        row = table.add_row().cells
        row[0].width = widths[0]
        row[1].width = widths[1]
        row[0].text = phase
        row[1].text = focus
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    set_run_font(run)

    document.add_paragraph("Suggested Solved Problem Targets", style="Heading 2")
    for label, target in TARGETS:
        p = document.add_paragraph(style="Body Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        set_run_font(run)
        p.add_run(target)


def add_topic_sheet(document):
    document.add_paragraph("Topic-Wise LeetCode Sheet", style="Heading 1")
    intro = document.add_paragraph(style="Normal")
    intro.add_run("Solve each topic in order. Finish beginner problems first, then intermediate, then advanced, and revise wrong problems after a few days.")

    for topic, levels in TOPIC_SHEET:
        document.add_paragraph(topic, style="Heading 2")
        for label in ("Beginner", "Intermediate", "Advanced"):
            add_topic_level(document, label, levels[label])

    document.add_paragraph("How to Use This Sheet", style="Heading 2")
    for item in (
        "Start with 4 to 6 problems per topic before moving deeper.",
        "Revisit wrong problems after 2 to 3 days.",
        "Maintain notes for patterns, mistakes, and reusable templates.",
        "Do not jump to dynamic programming too early; build comfort in arrays, strings, trees, and graphs first.",
    ):
        add_bullet(document, item)


def build_document():
    document = Document()
    apply_styles(document)
    add_footer(document)
    add_intro(document)
    add_syllabus_section(document, "Core DSA Syllabus", CORE_SYLLABUS)
    add_syllabus_section(document, "Advanced DSA Syllabus", ADVANCED_SYLLABUS)
    add_strategy_section(document)
    add_topic_sheet(document)
    document.save(OUTPUT_DOCX)


def html_list(items):
    return "".join(f"<li>{escape(item)}</li>" for item in items)


def html_syllabus_blocks(items):
    parts = []
    for section_title, bullets in items:
        parts.append(f"<section class='topic-block'><h3>{escape(section_title)}</h3><ul>{html_list(bullets)}</ul></section>")
    return "".join(parts)


def html_topic_sheet():
    blocks = []
    for topic, levels in TOPIC_SHEET:
        level_html = []
        for label in ("Beginner", "Intermediate", "Advanced"):
            joined = ", ".join(levels[label])
            level_html.append(
                f"<p class='level-line'><strong>{escape(label)}:</strong> {escape(joined)}</p>"
            )
        blocks.append(f"<section class='topic-block'><h3>{escape(topic)}</h3>{''.join(level_html)}</section>")
    return "".join(blocks)


def html_table_rows(rows):
    return "".join(
        f"<tr><td>{escape(left)}</td><td>{escape(right)}</td></tr>"
        for left, right in rows
    )


def build_html():
    html = f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>DSA Placement and LeetCode Guide</title>
  <style>
    @page {{
      size: A4;
      margin: 18mm 16mm 18mm 16mm;
    }}
    * {{
      box-sizing: border-box;
    }}
    body {{
      margin: 0;
      font-family: Arial, sans-serif;
      color: #111;
      background: #eef2f6;
    }}
    .page {{
      width: 100%;
      max-width: 920px;
      margin: 0 auto;
      padding: 28px 24px 40px;
      background: #ffffff;
    }}
    .hero {{
      border: 1px solid #d5dee9;
      background: linear-gradient(135deg, #f6f8fb 0%, #edf4fb 100%);
      border-radius: 18px;
      padding: 26px 28px;
      margin-bottom: 26px;
    }}
    h1 {{
      font-size: 28px;
      margin: 0 0 8px;
      line-height: 1.2;
    }}
    .subtitle {{
      margin: 0 0 8px;
      font-size: 14px;
      color: #44556b;
      line-height: 1.5;
    }}
    .note {{
      margin: 0;
      font-size: 13px;
      color: #5d6a7d;
      line-height: 1.45;
    }}
    h2 {{
      font-size: 22px;
      margin: 30px 0 14px;
      padding-bottom: 8px;
      border-bottom: 2px solid #e7edf4;
      page-break-after: avoid;
    }}
    h3 {{
      font-size: 16px;
      margin: 18px 0 8px;
      page-break-after: avoid;
    }}
    p, li {{
      font-size: 12.6px;
      line-height: 1.48;
    }}
    ul, ol {{
      margin: 8px 0 10px 18px;
      padding: 0;
    }}
    li {{
      margin: 0 0 6px;
    }}
    .topic-block {{
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .two-col {{
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 20px;
    }}
    .card {{
      border: 1px solid #dfe6ee;
      border-radius: 14px;
      padding: 14px 16px;
      background: #fbfcfe;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    .compact-list li {{
      margin-bottom: 4px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 8px 0 14px;
      font-size: 12.5px;
      break-inside: avoid;
      page-break-inside: avoid;
    }}
    th, td {{
      border: 1px solid #d9e1ea;
      padding: 9px 10px;
      text-align: left;
      vertical-align: top;
    }}
    th {{
      background: #f3f6fa;
      font-size: 12px;
    }}
    .level-line {{
      margin: 6px 0;
    }}
    .chips {{
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
      margin-top: 10px;
    }}
    .chip {{
      border: 1px solid #d9e2ec;
      background: #f7f9fc;
      border-radius: 999px;
      padding: 6px 10px;
      font-size: 11.6px;
    }}
    .footer-note {{
      margin-top: 22px;
      font-size: 11px;
      color: #66768a;
      text-align: center;
    }}
    @media print {{
      body {{
        background: #ffffff;
      }}
      .page {{
        max-width: none;
        padding: 0;
      }}
    }}
  </style>
</head>
<body>
  <main class="page">
    <section class="hero">
      <h1>DSA Placement and LeetCode Guide</h1>
      <p class="subtitle">A structured roadmap for placements, coding interviews, strong DSA fundamentals, and LeetCode practice.</p>
      <p class="note">This guide consolidates the full DSA syllabus, advanced DSA syllabus, best study order, highest-value placement topics, and a topic-wise LeetCode progression from beginner to advanced.</p>
    </section>

    <h2>Core DSA Syllabus</h2>
    {html_syllabus_blocks(CORE_SYLLABUS)}

    <h2>Advanced DSA Syllabus</h2>
    {html_syllabus_blocks(ADVANCED_SYLLABUS)}

    <h2>Placement Strategy</h2>
    <section class="two-col">
      <div class="card">
        <h3>Best Order to Study</h3>
        <ol class="compact-list">{html_list(BEST_ORDER)}</ol>
      </div>
      <div class="card">
        <h3>Highest-Value Topics for Placements</h3>
        <div class="chips">{"".join(f"<span class='chip'>{escape(item)}</span>" for item in KEY_FOCUS)}</div>
      </div>
    </section>

    <section class="topic-block">
      <h3>Practical Phases</h3>
      <table>
        <thead>
          <tr><th>Phase</th><th>Focus</th></tr>
        </thead>
        <tbody>
          {html_table_rows(PRACTICE_STAGES)}
        </tbody>
      </table>
    </section>

    <section class="topic-block">
      <h3>Suggested Solved Problem Targets</h3>
      <ul>{html_list([f"{label}: {target}" for label, target in TARGETS])}</ul>
    </section>

    <h2>Topic-Wise LeetCode Sheet</h2>
    <p>Solve each topic in order. Finish beginner problems first, then intermediate, then advanced, and revise wrong problems after a few days.</p>
    {html_topic_sheet()}

    <section class="topic-block">
      <h3>How to Use This Sheet</h3>
      <ul>
        {html_list([
            "Start with 4 to 6 problems per topic before moving deeper.",
            "Revisit wrong problems after 2 to 3 days.",
            "Maintain notes for patterns, mistakes, and reusable templates.",
            "Do not jump to dynamic programming too early; build comfort in arrays, strings, trees, and graphs first.",
        ])}
      </ul>
    </section>

    <p class="footer-note">Prepared for placement preparation, interview problem solving, and long-term DSA fundamentals.</p>
  </main>
</body>
</html>
"""
    with open(OUTPUT_HTML, "w", encoding="utf-8") as file:
        file.write(html)


if __name__ == "__main__":
    build_document()
    build_html()
